"""Input loaders: .docx / .md / .txt, plus direct text strings."""

from __future__ import annotations

from pathlib import Path


SUPPORTED_SUFFIXES = {".docx", ".md", ".txt", ".markdown"}


def load_document(path: str | Path) -> str:
    """Load full paper text from .docx, .md, or .txt."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Input file not found: {p}")
    suffix = p.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(
            f"Unsupported format '{suffix}'. Supported: .docx, .md, .txt"
        )
    if suffix == ".docx":
        return _load_docx(p)
    return p.read_text(encoding="utf-8")


def load_text(text: str) -> str:
    """Pass-through interface for direct string input."""
    if text is None:
        raise ValueError("text must not be None")
    return text.strip()


def _load_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for .docx input. "
            "Install with: pip install python-docx"
        ) from exc

    doc = Document(str(path))
    blocks: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            blocks.append(t)
    # Tables: keep cell text as lines
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n\n".join(blocks)
